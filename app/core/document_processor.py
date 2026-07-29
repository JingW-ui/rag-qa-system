# -*- coding: utf-8 -*-
"""
文档处理器 — 按文件类型解析 + 分块编排。
"""

import os
import json
from typing import Optional

from app.utils.chunker import chunk_text
from app.utils.text_cleaner import clean_text


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".markdown", ".json", ".jsonl"}


class DocumentProcessor:
    """解析 PDF/DOCX/MD/TXT/JSON/JSONL，返回纯文本，支持分块。"""

    def __init__(
        self,
        chunk_size: int = 800,
        chunk_overlap: int = 120,
        chunk_method: str = "recursive",
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunk_method = chunk_method

    # ------------------------------------------------------------------ #
    #  Public API
    # ------------------------------------------------------------------ #

    def parse(self, file_path: str) -> str:
        """
        根据扩展名分发解析。

        Returns:
            清洗后的纯文本。
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {ext}，支持: {SUPPORTED_EXTENSIONS}")

        if ext == ".pdf":
            raw = self._parse_pdf(file_path)
        elif ext in (".docx",):
            raw = self._parse_docx(file_path)
        elif ext in (".md", ".markdown", ".txt"):
            raw = self._parse_text(file_path)
        elif ext == ".json":
            raw = self._parse_json(file_path)
        elif ext == ".jsonl":
            raw = self._parse_jsonl(file_path)
        else:
            raise ValueError(f"未实现的解析器: {ext}")

        return clean_text(raw)

    def parse_and_chunk(self, file_path: str) -> tuple[str, list[str]]:
        """一步完成解析 + 分块。

        Returns:
            (全文, 分块列表)
        """
        full_text = self.parse(file_path)
        chunks = self.chunk(full_text)
        return full_text, chunks

    def chunk(self, text: str) -> list[str]:
        """对文本分块。"""
        return chunk_text(
            text,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            method=self.chunk_method,
        )

    # ------------------------------------------------------------------ #
    #  Internal parsers
    # ------------------------------------------------------------------ #

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        import fitz  # PyMuPDF (lazy import 避免启动加载)

        doc = fitz.open(file_path)
        parts: list[str] = []
        for page in doc:
            text = page.get_text("text")
            if text:
                parts.append(text)
        doc.close()
        return "\n\n".join(parts)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    @staticmethod
    def _parse_text(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    # ------------------------------------------------------------------ #
    #  JSON / JSONL 解析
    # ------------------------------------------------------------------ #

    @staticmethod
    def _parse_json(file_path: str) -> str:
        """解析 JSON 文件 — 自动判断 Array / Object 模式。

        - Array [...]  → 逐条扁平化（结构化数据模式）
        - Object {...} → pretty-print（配置/文档模式）
        """
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, list):
            return DocumentProcessor._flatten_json_array(data)
        elif isinstance(data, dict):
            return DocumentProcessor._flatten_json_object(data)
        else:
            # 基础类型（string/number/bool）直接转字符串
            return str(data)

    @staticmethod
    def _parse_jsonl(file_path: str) -> str:
        """解析 JSONL 文件 — 每行一个 JSON 对象，逐条扁平化。"""
        lines: list[str] = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    lines.append(DocumentProcessor._obj_to_flat_text(obj))
                except json.JSONDecodeError:
                    continue  # 跳过无效行
        return "\n".join(lines)

    # ---- JSON 转换辅助 ----

    @staticmethod
    def _flatten_json_array(data: list) -> str:
        """Array 模式：每条记录转成一行扁平文本。"""
        lines: list[str] = []
        for item in data:
            lines.append(DocumentProcessor._obj_to_flat_text(item))
        return "\n".join(lines)

    @staticmethod
    def _flatten_json_object(data: dict, indent: int = 0) -> str:
        """Object 模式：保留层级结构的 pretty-print。"""
        return DocumentProcessor._dict_to_indented(data, indent=0)

    @staticmethod
    def _obj_to_flat_text(obj) -> str:
        """将单个对象扁平化为 'key: value | key: value' 格式。"""
        if isinstance(obj, dict):
            parts: list[str] = []
            for k, v in obj.items():
                val_str = DocumentProcessor._value_to_str(v)
                parts.append(f"{k}: {val_str}")
            return " | ".join(parts)
        elif isinstance(obj, list):
            return ", ".join(DocumentProcessor._value_to_str(x) for x in obj)
        else:
            return str(obj)

    @staticmethod
    def _dict_to_indented(data: dict, indent: int = 0) -> str:
        """递归缩进输出嵌套 dict。"""
        prefix = "  " * indent
        lines: list[str] = []
        for k, v in data.items():
            if isinstance(v, dict):
                lines.append(f"{prefix}{k}:")
                lines.append(DocumentProcessor._dict_to_indented(v, indent + 1))
            elif isinstance(v, list):
                lines.append(f"{prefix}{k}:")
                for item in v:
                    lines.append(f"{prefix}  - {DocumentProcessor._value_to_str(item)}")
            else:
                lines.append(f"{prefix}{k}: {DocumentProcessor._value_to_str(v)}")
        return "\n".join(lines)

    @staticmethod
    def _value_to_str(v) -> str:
        """将值转为紧凑字符串。"""
        if v is None:
            return ""
        if isinstance(v, bool):
            return str(v).lower()
        if isinstance(v, (int, float)):
            return str(v)
        if isinstance(v, str):
            return v
        if isinstance(v, list):
            return ", ".join(DocumentProcessor._value_to_str(x) for x in v)
        if isinstance(v, dict):
            # 嵌套 dict 在扁平模式下用 JSON 紧凑表示
            return json.dumps(v, ensure_ascii=False, separators=(",", ":"))
        return str(v)
